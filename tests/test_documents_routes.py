"""`routes/documents.py` — cùng convention gọi thẳng hàm route + set `ContextVar` tay đã dùng ở
`test_sections_routes.py` (chưa có tiền lệ `TestClient` trong repo cho loại test này)."""

from __future__ import annotations

from collections.abc import AsyncIterator
from contextlib import asynccontextmanager
from io import BytesIO
from uuid import UUID, uuid4

import pytest
import pytest_asyncio
from fastapi import HTTPException, UploadFile
from studio_app import middleware
from studio_app.core._db import Pool, close_pools, get_pool
from studio_app.routes import documents as documents_module
from studio_app.routes.documents import upload_document
from studio_workbench.tenant_wall import ResolvedContext


@pytest_asyncio.fixture(autouse=True)
async def _close_singleton_pools_after_test() -> AsyncIterator[None]:
    yield
    await close_pools()


def _set_session(*, tenant_id: UUID, user: str, system_roles: list[str]) -> object:
    session = ResolvedContext(tenant_id=tenant_id, user=user, system_roles=system_roles)
    return middleware._request_session.set(session)


@asynccontextmanager
async def _simulate_request_connection() -> AsyncIterator[None]:
    pool = await get_pool()
    async with pool.connection() as conn:
        token = middleware._request_conn.set(conn)
        try:
            yield
        finally:
            middleware._request_conn.reset(token)


async def _seed_tenant(admin_pool: Pool, name: str) -> UUID:
    async with admin_pool.connection() as conn:
        cur = await conn.execute("INSERT INTO core.tenants (name) VALUES (%s) RETURNING id", (name,))
        row = await cur.fetchone()
    assert row is not None
    return UUID(str(row[0]))


async def _seed_user(admin_pool: Pool, tenant_id: UUID, email: str, system_roles: list[str]) -> UUID:
    async with admin_pool.connection() as conn:
        cur = await conn.execute(
            "INSERT INTO core.users (tenant_id, email, password_hash, system_roles) "
            "VALUES (%s, %s, %s, %s) RETURNING id",
            (str(tenant_id), email, "not-a-real-hash", system_roles),
        )
        row = await cur.fetchone()
    assert row is not None
    return UUID(str(row[0]))


async def _seed_section(admin_pool: Pool, tenant_id: UUID, name: str, created_by: UUID) -> None:
    async with admin_pool.connection() as conn:
        await conn.execute(
            "INSERT INTO core.sections (tenant_id, name, created_by) VALUES (%s, %s, %s)",
            (str(tenant_id), name, str(created_by)),
        )


def _md_upload_file(filename: str, content: bytes) -> UploadFile:
    return UploadFile(filename=filename, file=BytesIO(content))


_VALID_MD = b"## Nghi phep\nBao truoc 3 ngay lam viec.\n"


async def test_company_admin_uploads_document(admin_pool: Pool) -> None:
    tenant_id = await _seed_tenant(admin_pool, "documents-probe-a")
    admin_id = await _seed_user(admin_pool, tenant_id, "admin@acme.com", ["admin"])
    await _seed_section(admin_pool, tenant_id, "hr", admin_id)

    token = _set_session(tenant_id=tenant_id, user="admin@acme.com", system_roles=["admin"])
    try:
        async with _simulate_request_connection():
            result = await upload_document(
                file=_md_upload_file("leave.md", _VALID_MD),
                section_role="hr",
                tenant_id=None,
            )
    finally:
        middleware._request_session.reset(token)  # type: ignore[arg-type]

    assert result.section_role == "hr"
    assert result.chunk_count == 1
    # `doc_id` (cột `kb.chunks.doc_id`, tách khỏi PK `chunk_id`) = `{slug(section_role)}-{slug(tên
    # file)}-{đuôi file}`, KHÔNG còn tenant-hex-prefixed — xem docstring `routes/documents.py`
    # (quyết định doc_id column, review app#58; đuôi file thêm sau, phát hiện thật trên tenant
    # Ricons — xem `test_hai_duoi_file_cung_ten_khong_ghi_de_nhau`).
    assert result.doc_id == "hr-leave-md"

    async with admin_pool.connection() as conn:
        # `kb.chunks` có `FORCE ROW LEVEL SECURITY` — cắn cả owner nếu chưa `set_config`
        # `app.tenant_id` (xem `postgres.py::_bind_tenant`), nên phải bind trước khi verify.
        await conn.execute("SELECT set_config('app.tenant_id', %s, true)", (str(tenant_id),))
        cur = await conn.execute("SELECT count(*) FROM kb.chunks WHERE tenant_id = %s", (str(tenant_id),))
        row = await cur.fetchone()
    assert row is not None
    assert row[0] == 1


async def test_upload_doc_id_la_slug_ten_file(admin_pool: Pool) -> None:
    """`doc_id` = `{slug(section_role)}-{slug(stem)}-{đuôi file}` — hoa/khoảng trắng trong tên
    file phải bị gộp/hạ chữ (khớp quyết định doc_id column, khác `chunk_id_prefix` vẫn giữ hash để
    tránh đụng PK)."""
    tenant_id = await _seed_tenant(admin_pool, "documents-probe-m")
    admin_id = await _seed_user(admin_pool, tenant_id, "admin@acme.com", ["admin"])
    await _seed_section(admin_pool, tenant_id, "hr", admin_id)

    token = _set_session(tenant_id=tenant_id, user="admin@acme.com", system_roles=["admin"])
    try:
        async with _simulate_request_connection():
            result = await upload_document(
                file=_md_upload_file("Bao Cao Q1.md", _VALID_MD),
                section_role="hr",
                tenant_id=None,
            )
    finally:
        middleware._request_session.reset(token)  # type: ignore[arg-type]

    assert result.doc_id == "hr-bao-cao-q1-md"
    # `doc_name` KHÔNG qua `_slugify` — giữ nguyên hoa/thường/khoảng trắng của tên file gốc, khác
    # hẳn `doc_id`. Đây là điểm luật hiển thị đòi hỏi: `doc_id` mất thông tin (chữ hoa, khoảng
    # trắng) nên không dùng được để hiện lên UI.
    assert result.doc_name == "Bao Cao Q1"


async def test_upload_luu_doc_name_vao_kb_chunks(admin_pool: Pool) -> None:
    """`doc_name` phải lên CẢ cột `kb.chunks.doc_name`, không chỉ nằm trong response — nếu không
    thì mọi đường đọc lại sau (list tài liệu, `chunks_for_tenant`) mất nhãn hiển thị, chỉ còn
    `doc_id` (slug, cấm hiển thị thẳng lên UI). Giữ dấu tiếng Việt nguyên vẹn — đúng test chống
    rỗng-nghĩa cho luật hiển thị."""
    tenant_id = await _seed_tenant(admin_pool, "documents-probe-p")
    admin_id = await _seed_user(admin_pool, tenant_id, "admin@acme.com", ["admin"])
    await _seed_section(admin_pool, tenant_id, "hr", admin_id)

    token = _set_session(tenant_id=tenant_id, user="admin@acme.com", system_roles=["admin"])
    try:
        async with _simulate_request_connection():
            result = await upload_document(
                file=_md_upload_file("Chế độ nghỉ phép.md", _VALID_MD),
                section_role="hr",
                tenant_id=None,
            )
    finally:
        middleware._request_session.reset(token)  # type: ignore[arg-type]

    assert result.doc_name == "Chế độ nghỉ phép"

    async with admin_pool.connection() as conn:
        await conn.execute("SELECT set_config('app.tenant_id', %s, true)", (str(tenant_id),))
        cur = await conn.execute("SELECT DISTINCT doc_name FROM kb.chunks WHERE tenant_id = %s", (str(tenant_id),))
        rows = await cur.fetchall()
    assert rows == [("Chế độ nghỉ phép",)]


async def test_reupload_ten_trung_khac_phong_ban_khong_xoa_nham(admin_pool: Pool) -> None:
    """review app#58 (dholmes0207) — tái hiện đúng ca bug: 2 tài liệu KHÁC NHAU, cùng tên file,
    khác `section_role`. Bản trước khi vá dùng `doc_id = slug(stem)` (không mang phòng ban), nên
    upload file thứ hai xoá LUÔN tài liệu phòng ban đầu qua `delete_by_doc_id` (không lọc phòng
    ban). `section_role` trong khoá `doc_id` là ranh giới phân quyền đọc thật
    (`fetch_tenant_section_names`), khác hẳn một re-upload cùng tài liệu."""
    tenant_id = await _seed_tenant(admin_pool, "documents-probe-o")
    admin_id = await _seed_user(admin_pool, tenant_id, "admin@acme.com", ["admin"])
    await _seed_section(admin_pool, tenant_id, "hr", admin_id)
    await _seed_section(admin_pool, tenant_id, "finance", admin_id)

    token = _set_session(tenant_id=tenant_id, user="admin@acme.com", system_roles=["admin"])
    try:
        async with _simulate_request_connection():
            hr_result = await upload_document(
                file=_md_upload_file("leave.md", b"## Nghi phep\nQuy dinh nghi phep phong nhan su."),
                section_role="hr",
                tenant_id=None,
            )
            finance_result = await upload_document(
                file=_md_upload_file("leave.md", b"## Han muc\nQuy dinh han muc chi tieu phong tai chinh."),
                section_role="finance",
                tenant_id=None,
            )
    finally:
        middleware._request_session.reset(token)  # type: ignore[arg-type]

    assert hr_result.doc_id != finance_result.doc_id

    async with admin_pool.connection() as conn:
        await conn.execute("SELECT set_config('app.tenant_id', %s, true)", (str(tenant_id),))
        cur = await conn.execute(
            "SELECT section_role, count(*) FROM kb.chunks WHERE tenant_id = %s GROUP BY section_role",
            (str(tenant_id),),
        )
        rows = await cur.fetchall()
    assert dict(rows) == {"hr": 1, "finance": 1}, "cả hai tài liệu phải còn sống — không cái nào bị xoá nhầm"


async def test_hai_duoi_file_cung_ten_khong_ghi_de_nhau(admin_pool: Pool) -> None:
    """2 file GỐC khác nhau, CÙNG `section_role`, CÙNG stem (`Path.stem` bỏ đuôi trước khi vào
    `doc_id`), khác đuôi (`.md` vs `.txt`) — trước bản vá này, `doc_id` không mang đuôi file nên
    2 file đụng `doc_id`, file sau ghi đè êm file trước qua `delete_by_doc_id` (tái hiện đúng ca
    thật: `quy chế lương thưởng - hr.docx` và `quy chế lương thưởng - hr.md` cùng tồn tại trên đĩa
    nhưng chỉ 1 trong 2 sống sót trong `kb.chunks`). `doc_id` giờ mang thêm đuôi file
    (`{slug(role)}-{slug(stem)}-{đuôi}`) nên 2 file này phải là 2 tài liệu ĐỘC LẬP, không cái nào
    xoá cái kia."""
    tenant_id = await _seed_tenant(admin_pool, "documents-probe-q")
    admin_id = await _seed_user(admin_pool, tenant_id, "admin@acme.com", ["admin"])
    await _seed_section(admin_pool, tenant_id, "hr", admin_id)

    token = _set_session(tenant_id=tenant_id, user="admin@acme.com", system_roles=["admin"])
    try:
        async with _simulate_request_connection():
            md_result = await upload_document(
                file=_md_upload_file("quy che luong thuong.md", b"## Muc luong\nQuy dinh muc luong co ban."),
                section_role="hr",
                tenant_id=None,
            )
            txt_result = await upload_document(
                file=_md_upload_file("quy che luong thuong.txt", b"Quy dinh muc thuong quy cua cong ty."),
                section_role="hr",
                tenant_id=None,
            )
    finally:
        middleware._request_session.reset(token)  # type: ignore[arg-type]

    assert md_result.doc_id != txt_result.doc_id

    async with admin_pool.connection() as conn:
        await conn.execute("SELECT set_config('app.tenant_id', %s, true)", (str(tenant_id),))
        cur = await conn.execute(
            "SELECT doc_id, count(*) FROM kb.chunks WHERE tenant_id = %s GROUP BY doc_id",
            (str(tenant_id),),
        )
        rows = dict(await cur.fetchall())
    assert rows == {md_result.doc_id: 1, txt_result.doc_id: 1}, "cả hai đuôi file phải còn sống — không cái nào bị đè"


async def test_reupload_cung_doc_id_xoa_orphan_chunk(admin_pool: Pool) -> None:
    """Re-upload CÙNG tên file với nội dung NGẮN HƠN (ít chunk hơn) không còn để lại `chunk_id`
    mồ côi — đóng giới hạn cũ đã ghi trong docstring `routes/documents.py` (route giờ gọi
    `pipeline.delete_by_doc_id` trước khi ghi chunk mới)."""
    tenant_id = await _seed_tenant(admin_pool, "documents-probe-n")
    admin_id = await _seed_user(admin_pool, tenant_id, "admin@acme.com", ["admin"])
    await _seed_section(admin_pool, tenant_id, "hr", admin_id)

    # 1000 từ > WORDS_PER_CHUNK (850) → 2 chunk ở lần upload đầu.
    long_text = " ".join(f"tu{i}" for i in range(1000)).encode()
    short_text = b"noi dung rat ngan sau khi sua lai file."

    token = _set_session(tenant_id=tenant_id, user="admin@acme.com", system_roles=["admin"])
    try:
        async with _simulate_request_connection():
            first = await upload_document(
                file=_md_upload_file("bao-cao.txt", long_text),
                section_role="hr",
                tenant_id=None,
            )
            assert first.chunk_count == 2

            second = await upload_document(
                file=_md_upload_file("bao-cao.txt", short_text),
                section_role="hr",
                tenant_id=None,
            )
    finally:
        middleware._request_session.reset(token)  # type: ignore[arg-type]

    assert second.chunk_count == 1
    assert second.doc_id == first.doc_id

    async with admin_pool.connection() as conn:
        await conn.execute("SELECT set_config('app.tenant_id', %s, true)", (str(tenant_id),))
        cur = await conn.execute(
            "SELECT count(*) FROM kb.chunks WHERE tenant_id = %s AND doc_id = %s",
            (str(tenant_id), second.doc_id),
        )
        row = await cur.fetchone()
    assert row is not None
    assert row[0] == 1  # không còn 1 chunk mồ côi từ lần upload trước (2 - 1 = 1 dư nếu chưa vá)


async def test_upload_rejects_unknown_section_role(admin_pool: Pool) -> None:
    tenant_id = await _seed_tenant(admin_pool, "documents-probe-b")
    await _seed_user(admin_pool, tenant_id, "admin@acme.com", ["admin"])

    token = _set_session(tenant_id=tenant_id, user="admin@acme.com", system_roles=["admin"])
    try:
        with pytest.raises(HTTPException) as exc_info:
            async with _simulate_request_connection():
                await upload_document(
                    file=_md_upload_file("leave.md", _VALID_MD),
                    section_role="not-a-real-department",
                    tenant_id=None,
                )
        assert exc_info.value.status_code == 400
    finally:
        middleware._request_session.reset(token)  # type: ignore[arg-type]


async def test_upload_rejects_unsupported_extension(admin_pool: Pool) -> None:
    """`.pdf` (ngoài `extract.SUPPORTED_SUFFIXES`) vẫn bị chặn — khác `.txt` (đã CHO PHÉP từ khi
    route chuyển sang `chunk_window.cut_window`, xem `test_upload_accepts_txt_file` bên dưới)."""
    tenant_id = await _seed_tenant(admin_pool, "documents-probe-c")
    admin_id = await _seed_user(admin_pool, tenant_id, "admin@acme.com", ["admin"])
    await _seed_section(admin_pool, tenant_id, "hr", admin_id)

    token = _set_session(tenant_id=tenant_id, user="admin@acme.com", system_roles=["admin"])
    try:
        with pytest.raises(HTTPException) as exc_info:
            async with _simulate_request_connection():
                await upload_document(
                    file=_md_upload_file("leave.pdf", _VALID_MD),
                    section_role="hr",
                    tenant_id=None,
                )
        assert exc_info.value.status_code == 422
    finally:
        middleware._request_session.reset(token)  # type: ignore[arg-type]


async def test_upload_accepts_txt_file(admin_pool: Pool) -> None:
    """`.txt` — trước đây bị 422 ngay ở kiểm đuôi file (khi route chỉ nhận `.md`); giờ đi qua
    `extract.extract_text` + `chunk_window.cut_window`, không đòi hỏi cấu trúc heading `##`."""
    tenant_id = await _seed_tenant(admin_pool, "documents-probe-h")
    admin_id = await _seed_user(admin_pool, tenant_id, "admin@acme.com", ["admin"])
    await _seed_section(admin_pool, tenant_id, "hr", admin_id)

    token = _set_session(tenant_id=tenant_id, user="admin@acme.com", system_roles=["admin"])
    try:
        async with _simulate_request_connection():
            result = await upload_document(
                file=_md_upload_file("ghi-chu.txt", b"Ghi chu roi rac, khong co heading nao ca."),
                section_role="hr",
                tenant_id=None,
            )
    finally:
        middleware._request_session.reset(token)  # type: ignore[arg-type]

    assert result.chunk_count == 1


async def test_upload_accepts_md_without_heading(admin_pool: Pool) -> None:
    """Trước đây (`_cut_document`, cắt theo heading `##`) file `.md` không heading bị 422. Giờ route
    dùng `chunk_window.cut_window` (không đòi cấu trúc), nên phải THÀNH CÔNG — đúng động lực chính
    của thay đổi này: ghi chú `.md` thật thường không có heading chuẩn."""
    tenant_id = await _seed_tenant(admin_pool, "documents-probe-d")
    admin_id = await _seed_user(admin_pool, tenant_id, "admin@acme.com", ["admin"])
    await _seed_section(admin_pool, tenant_id, "hr", admin_id)

    token = _set_session(tenant_id=tenant_id, user="admin@acme.com", system_roles=["admin"])
    try:
        async with _simulate_request_connection():
            result = await upload_document(
                file=_md_upload_file("leave.md", b"khong co heading nao ca"),
                section_role="hr",
                tenant_id=None,
            )
    finally:
        middleware._request_session.reset(token)  # type: ignore[arg-type]

    assert result.chunk_count == 1


async def test_upload_accepts_docx_file(admin_pool: Pool) -> None:
    """`.docx` thật (dựng bằng `python-docx`, không phải giả `.md` đổi tên) — đi qua
    `extract._extract_docx` rồi `chunk_window.cut_window`, ghi thật vào `kb.chunks`."""
    from io import BytesIO as _BytesIO

    from docx import Document

    doc = Document()
    doc.add_paragraph("Chinh sach nghi phep.")
    doc.add_paragraph("Bao truoc it nhat 3 ngay lam viec.")
    buf = _BytesIO()
    doc.save(buf)

    tenant_id = await _seed_tenant(admin_pool, "documents-probe-i")
    admin_id = await _seed_user(admin_pool, tenant_id, "admin@acme.com", ["admin"])
    await _seed_section(admin_pool, tenant_id, "hr", admin_id)

    token = _set_session(tenant_id=tenant_id, user="admin@acme.com", system_roles=["admin"])
    try:
        async with _simulate_request_connection():
            result = await upload_document(
                file=_md_upload_file("chinh-sach.docx", buf.getvalue()),
                section_role="hr",
                tenant_id=None,
            )
    finally:
        middleware._request_session.reset(token)  # type: ignore[arg-type]

    assert result.chunk_count == 1

    async with admin_pool.connection() as conn:
        await conn.execute("SELECT set_config('app.tenant_id', %s, true)", (str(tenant_id),))
        cur = await conn.execute("SELECT text FROM kb.chunks WHERE tenant_id = %s", (str(tenant_id),))
        row = await cur.fetchone()
    assert row is not None
    assert "Bao truoc it nhat 3 ngay lam viec." in row[0]


async def test_company_admin_cannot_upload_for_other_tenant(admin_pool: Pool) -> None:
    tenant_a = await _seed_tenant(admin_pool, "documents-probe-e-a")
    tenant_b = await _seed_tenant(admin_pool, "documents-probe-e-b")
    admin_a_id = await _seed_user(admin_pool, tenant_a, "admin-a@acme.com", ["admin"])
    await _seed_section(admin_pool, tenant_a, "hr", admin_a_id)

    token = _set_session(tenant_id=tenant_a, user="admin-a@acme.com", system_roles=["admin"])
    try:
        with pytest.raises(HTTPException) as exc_info:
            async with _simulate_request_connection():
                await upload_document(
                    file=_md_upload_file("leave.md", _VALID_MD),
                    section_role="hr",
                    tenant_id=str(tenant_b),
                )
        assert exc_info.value.status_code == 403
    finally:
        middleware._request_session.reset(token)  # type: ignore[arg-type]


async def test_superadmin_must_declare_tenant_id(admin_pool: Pool) -> None:
    tenant_id = await _seed_tenant(admin_pool, "documents-probe-f")
    await _seed_user(admin_pool, tenant_id, "su@sys", ["superadmin"])

    token = _set_session(tenant_id=tenant_id, user="su@sys", system_roles=["superadmin"])
    try:
        with pytest.raises(HTTPException) as exc_info:
            async with _simulate_request_connection():
                await upload_document(
                    file=_md_upload_file("leave.md", _VALID_MD),
                    section_role="hr",
                    tenant_id=None,
                )
        assert exc_info.value.status_code == 400
    finally:
        middleware._request_session.reset(token)  # type: ignore[arg-type]


async def test_upload_rejects_unknown_tenant_for_superadmin(admin_pool: Pool) -> None:
    tenant_id = await _seed_tenant(admin_pool, "documents-probe-g")
    await _seed_user(admin_pool, tenant_id, "su@sys", ["superadmin"])

    token = _set_session(tenant_id=tenant_id, user="su@sys", system_roles=["superadmin"])
    try:
        with pytest.raises(HTTPException) as exc_info:
            async with _simulate_request_connection():
                await upload_document(
                    file=_md_upload_file("leave.md", _VALID_MD),
                    section_role="hr",
                    tenant_id=str(uuid4()),
                )
        assert exc_info.value.status_code == 404
    finally:
        middleware._request_session.reset(token)  # type: ignore[arg-type]


async def test_upload_rejects_document_over_word_cap(admin_pool: Pool, monkeypatch: pytest.MonkeyPatch) -> None:
    """`_MAX_UPLOAD_BYTES` chặn SỐ BYTE, không chặn KHỐI LƯỢNG CHỮ — và từ khi `.docx` vào danh
    sách thì hai thứ đó không còn tỉ lệ với nhau: `.docx` là file ZIP, nội dung được NÉN. Đo thật
    trên cùng hạn mức 1 MiB: `.txt` thuần ~168.000 từ (~247 chunk), `.docx` văn xuôi ~606.000 từ
    (~891 chunk), `.docx` nội dung lặp ~14.400.000 từ (~21.000 chunk). Mỗi chunk là một mục trong
    lô embedding ≤90 → tối đa ~234 lời gọi API tuần tự (timeout 120s/lô) nằm trong ĐÚNG MỘT
    request. `_MAX_WORDS` là cửa chặn theo đúng đơn vị sinh ra chi phí.

    Hạ hạn mức xuống 50 thay vì dựng fixture triệu từ: thứ được khoá ở đây là CÓ cửa chặn và nó
    trả 422, không phải con số 200.000 (số đó khoá riêng ở
    `test_word_cap_khong_siet_cua_dang_mo`)."""
    monkeypatch.setattr(documents_module, "_MAX_WORDS", 50)
    tenant_id = await _seed_tenant(admin_pool, "documents-probe-j")
    admin_id = await _seed_user(admin_pool, tenant_id, "admin@acme.com", ["admin"])
    await _seed_section(admin_pool, tenant_id, "hr", admin_id)

    token = _set_session(tenant_id=tenant_id, user="admin@acme.com", system_roles=["admin"])
    try:
        with pytest.raises(HTTPException) as exc_info:
            async with _simulate_request_connection():
                await upload_document(
                    file=_md_upload_file("dai.txt", " ".join(f"tu{i}" for i in range(51)).encode()),
                    section_role="hr",
                    tenant_id=None,
                )
        assert exc_info.value.status_code == 422
    finally:
        middleware._request_session.reset(token)  # type: ignore[arg-type]


async def test_upload_accepts_document_dung_bang_word_cap(admin_pool: Pool, monkeypatch: pytest.MonkeyPatch) -> None:
    """Vế đối xứng: ĐÚNG BẰNG hạn mức phải QUA. Không có test này thì đổi `>` thành `>=` (hoặc
    ngược lại) vẫn xanh — cửa chặn lệch một từ là loại lỗi im lặng nhất ở đây."""
    monkeypatch.setattr(documents_module, "_MAX_WORDS", 50)
    tenant_id = await _seed_tenant(admin_pool, "documents-probe-k")
    admin_id = await _seed_user(admin_pool, tenant_id, "admin@acme.com", ["admin"])
    await _seed_section(admin_pool, tenant_id, "hr", admin_id)

    token = _set_session(tenant_id=tenant_id, user="admin@acme.com", system_roles=["admin"])
    try:
        async with _simulate_request_connection():
            result = await upload_document(
                file=_md_upload_file("vua.txt", " ".join(f"tu{i}" for i in range(50)).encode()),
                section_role="hr",
                tenant_id=None,
            )
    finally:
        middleware._request_session.reset(token)  # type: ignore[arg-type]

    assert result.chunk_count == 1


def test_word_cap_khong_siet_cua_dang_mo() -> None:
    """`_MAX_WORDS` phải RỘNG HƠN thứ 1 MiB văn bản thuần từng cho qua, nếu không bản vá này biến
    thành một thoái lui: `.md`/`.txt` hợp lệ đang upload được hôm nay sẽ bắt đầu ăn 422. Cận trên
    của `.txt` 1 MiB là ~524.000 từ (từ 1 ký tự + 1 dấu cách = 2 byte/từ) ở ca suy biến, ~168.000
    ở văn xuôi thật — chọn 200.000 nằm trên mức văn xuôi thật, và ca suy biến 1-ký-tự-mỗi-từ
    không phải tài liệu người ta upload."""
    assert documents_module._MAX_WORDS >= 200_000


def test_slugify_ha_dau_tieng_viet_ve_chu_cai_goc() -> None:
    """`_slugify` trước bản vá chỉ giữ `[a-z0-9]`, nên MỖI nguyên âm có dấu (nằm ngoài tập đó) bị
    gộp chung với khoảng trắng xung quanh thành một dấu `-` — "Nghỉ phép" ra `"ngh-ph-p"` chứ
    không phải `"nghi-phep"`. Bản vá phải hạ dấu về chữ cái gốc TRƯỚC khi lọc, để mỗi âm tiết vẫn
    còn nguyên chữ cái của nó, đúng cách một slug tiếng Việt nên đọc được ngược lại là từ gì."""
    assert documents_module._slugify("Nghỉ phép") == "nghi-phep"
    # "đ" (U+0111) không có decomposition NFKD — không tự rã ra "d" + dấu như các nguyên âm khác,
    # nên cần xử lý riêng, không thì bị `.encode("ascii", "ignore")` xoá mất luôn cả phụ âm.
    assert documents_module._slugify("Chế độ nghỉ phép") == "che-do-nghi-phep"
    assert documents_module._slugify("Đà Nẵng") == "da-nang"


async def test_upload_rejects_corrupt_docx_bang_422_khong_phai_500(admin_pool: Pool) -> None:
    """Đuôi `.docx` nhưng ruột không phải gói OOXML — thao tác thật: bị báo "chỉ nhận
    .md/.txt/.docx" rồi đổi thẳng đuôi `.doc`/`.pdf` thành `.docx`.

    Route chỉ bắt `UnsupportedFormatError`, mà `python-docx` ném ra `zipfile.BadZipFile`/`KeyError`/
    `lxml.etree.XMLSyntaxError` — KHÔNG cái nào là `ValueError`. Trước bản vá ở
    `studio_kb.extract._extract_docx`, cả ba thoát nguyên dạng ra ngoài và thành HTTP 500. Test này
    nằm ở app (không phải kb) vì thứ cần khoá là MẶT TIẾP GIÁP: kb quy lỗi về `UnsupportedFormatError`
    còn route dịch nó sang 422 — hỏng một trong hai vế là quay lại 500, và CI của mỗi repo con
    không nhìn thấy vế bên kia."""
    tenant_id = await _seed_tenant(admin_pool, "documents-probe-l")
    admin_id = await _seed_user(admin_pool, tenant_id, "admin@acme.com", ["admin"])
    await _seed_section(admin_pool, tenant_id, "hr", admin_id)

    token = _set_session(tenant_id=tenant_id, user="admin@acme.com", system_roles=["admin"])
    try:
        with pytest.raises(HTTPException) as exc_info:
            async with _simulate_request_connection():
                await upload_document(
                    file=_md_upload_file("chinh-sach.docx", b"# Chinh sach nghi phep\nnoi dung"),
                    section_role="hr",
                    tenant_id=None,
                )
        assert exc_info.value.status_code == 422
    finally:
        middleware._request_session.reset(token)  # type: ignore[arg-type]
